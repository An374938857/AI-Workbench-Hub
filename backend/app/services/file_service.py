from datetime import date, datetime
from typing import Any

import chardet


def extract_text(file_path: str, file_type: str) -> str:
    ext = file_type.lower()

    if ext in ("txt", "md"):
        return _read_text_file(file_path)
    elif ext == "pdf":
        return _extract_pdf(file_path)
    elif ext == "docx":
        return _extract_docx(file_path)
    elif ext == "xlsx":
        return _extract_xlsx(file_path)
    elif ext == "csv":
        return _extract_csv(file_path)
    else:
        return ""


def _read_text_file(path: str) -> str:
    with open(path, "rb") as f:
        raw = f.read()
    detected = chardet.detect(raw)
    encoding = detected.get("encoding") or "utf-8"
    try:
        return raw.decode(encoding)
    except (UnicodeDecodeError, LookupError):
        return raw.decode("utf-8", errors="replace")


def _extract_docx(path: str) -> str:
    from docx import Document

    doc = Document(path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_pdf(path: str) -> str:
    # 首选 pypdf，失败时尝试 pdfplumber
    try:
        from pypdf import PdfReader

        reader = PdfReader(path)
        parts: list[str] = []
        for page in reader.pages:
            text = page.extract_text() or ""
            if text.strip():
                parts.append(text)
        return "\n\n".join(parts)
    except Exception:
        try:
            import pdfplumber

            parts: list[str] = []
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text() or ""
                    if text.strip():
                        parts.append(text)
            return "\n\n".join(parts)
        except Exception:
            return ""


def _extract_xlsx(path: str) -> str:
    import pandas as pd

    try:
        xls = pd.ExcelFile(path, engine="openpyxl")
        parts = []
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name)
            parts.append(f"[Sheet: {sheet_name}]")
            parts.append(df.to_string(index=False, max_rows=200))
        return "\n\n".join(parts)
    except Exception as e:
        # 兼容“扩展名是 xlsx，但实际是 csv/文本”的场景
        try:
            csv_text = _extract_csv(path)
            return (
                "[提示] 文件扩展名为 .xlsx，但内容不是标准 Excel 工作簿，已按 CSV 文本方式尝试预览。\n\n"
                + csv_text
            )
        except Exception:
            return f"[提示] 无法解析为标准 XLSX 文件：{str(e)}。请下载后使用本地软件打开。"


def _normalize_xlsx_cell(value: Any) -> Any:
    if value is None:
        return None

    try:
        import pandas as pd

        if pd.isna(value):
            return None
    except Exception:
        pass

    if hasattr(value, "item"):
        try:
            value = value.item()
        except Exception:
            pass

    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d %H:%M:%S")
    if isinstance(value, date):
        return value.isoformat()
    if isinstance(value, (str, int, float, bool)):
        return value

    return str(value)


def extract_xlsx_table(path: str, max_rows: int = 200, max_cols: int = 50) -> dict[str, Any]:
    import pandas as pd

    xls = pd.ExcelFile(path, engine="openpyxl")
    sheets: list[dict[str, Any]] = []

    for sheet_name in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet_name, dtype=object)
        total_rows = len(df)
        total_columns = len(df.columns)

        limited_df = df.iloc[:max_rows, :max_cols]
        columns = [str(col) if col is not None else "" for col in limited_df.columns.tolist()]

        rows: list[dict[str, Any]] = []
        for record in limited_df.to_dict(orient="records"):
            rows.append(
                {
                    (str(key) if key is not None else ""): _normalize_xlsx_cell(value)
                    for key, value in record.items()
                }
            )

        sheets.append(
            {
                "name": sheet_name,
                "columns": columns,
                "rows": rows,
                "total_rows": total_rows,
                "total_columns": total_columns,
                "truncated_rows": total_rows > max_rows,
                "truncated_columns": total_columns > max_cols,
            }
        )

    return {"sheets": sheets}


def _extract_csv(path: str) -> str:
    import pandas as pd

    with open(path, "rb") as f:
        raw = f.read(4096)
    detected = chardet.detect(raw)
    encoding = detected.get("encoding") or "utf-8"

    try:
        df = pd.read_csv(path, encoding=encoding, nrows=200)
    except Exception:
        df = pd.read_csv(path, encoding="utf-8", errors="replace", nrows=200)

    return df.to_string(index=False)
